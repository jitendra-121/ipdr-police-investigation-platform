import pandas as pd
import json
from pathlib import Path
import sys
from datetime import datetime, date, time
import numpy as np
from schema_mapper import SchemaMapper
from supabase_handler import SupabaseHandler
from typing import List, Dict

class FileProcessor:
    def __init__(self, verbose=False):
        self.supported_types = ['csv', 'excel', 'word']
        self.verbose = verbose
        self.schema_mapper = SchemaMapper()
        self.supabase_handler = SupabaseHandler(verbose=verbose)
        
    def _clean_data_for_json(self, data):
        """Clean data to ensure JSON serialization compatibility"""
        if isinstance(data, list):
            return [self._clean_record(record) for record in data]
        elif isinstance(data, dict):
            return {key: self._clean_data_for_json(value) for key, value in data.items()}
        else:
            return self._clean_value(data)
    
    def _clean_record(self, record):
        """Clean a single record/dictionary"""
        if not isinstance(record, dict):
            return record
        
        cleaned_record = {}
        for key, value in record.items():
            cleaned_record[key] = self._clean_value(value)
        
        return cleaned_record
    
    def _clean_value(self, value):
        """Clean individual values for JSON compatibility"""
        # Handle pandas NaN/None values
        if pd.isna(value) or value is None:
            return None
        
        # Handle datetime objects
        elif isinstance(value, (datetime, pd.Timestamp)):
            return value.strftime('%Y-%m-%d %H:%M:%S')
        
        # Handle date objects
        elif isinstance(value, date):
            return value.strftime('%Y-%m-%d')
        
        # Handle time objects
        elif isinstance(value, time):
            return value.strftime('%H:%M:%S')
        
        # Handle numpy types
        elif hasattr(np, 'integer') and isinstance(value, np.integer):
            return int(value)
        elif hasattr(np, 'floating') and isinstance(value, np.floating):
            return float(value)
        elif hasattr(np, 'bool_') and isinstance(value, np.bool_):
            return bool(value)
        elif isinstance(value, np.ndarray):
            return value.tolist()
        
        # Handle pandas Timestamp specifically
        elif hasattr(value, 'isoformat'):
            return value.isoformat()
        
        # Return as-is for basic types
        else:
            return value
        
    def process_csv(self, file_path):
        """Process CSV file and return data as JSON-compatible dict/list"""
        try:
            # Read CSV file
            df = pd.read_csv(file_path)
            
            # Convert to list of dictionaries (records format)
            data = df.to_dict('records')
            
            # Clean all data for JSON compatibility
            cleaned_data = self._clean_data_for_json(data)
            
            return cleaned_data
            
        except Exception as e:
            print(f"Error processing CSV file: {e}")
            return None
    
    def process_excel(self, file_path):
        """Process Excel file and return data as JSON-compatible dict/list"""
        try:
            # Read Excel file - get all sheet names first
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            
            # If only one sheet, return list of records
            if len(sheet_names) == 1:
                df = pd.read_excel(file_path, sheet_name=sheet_names[0])
                data = df.to_dict('records')
                
                # Clean all data for JSON compatibility
                cleaned_data = self._clean_data_for_json(data)
                
                return cleaned_data
            
            # If multiple sheets, return dict with sheet names as keys
            else:
                all_data = {}
                
                for sheet in sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet)
                    sheet_data = df.to_dict('records')
                    
                    # Clean all data for JSON compatibility
                    cleaned_data = self._clean_data_for_json(sheet_data)
                    
                    all_data[sheet] = cleaned_data
                
                return all_data
                
        except Exception as e:
            print(f"Error processing Excel file: {e}")
            return None
    
    def process_word(self, file_path):
        """Process Word document and return text/tables as JSON-compatible dict"""
        try:
            # Check if python-docx is installed
            try:
                from docx import Document
            except ImportError:
                print("python-docx not installed. Installing...")
                import subprocess
                subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
                from docx import Document
            
            doc = Document(file_path)
            
            # Extract all text
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract tables if any
            tables_data = []
            for table in doc.tables:
                table_content = []
                
                # Get headers from first row
                headers = []
                if len(table.rows) > 0:
                    for cell in table.rows[0].cells:
                        headers.append(cell.text.strip())
                
                # Get data from remaining rows
                for row_idx, row in enumerate(table.rows[1:], 1):
                    row_data = {}
                    for col_idx, cell in enumerate(row.cells):
                        if col_idx < len(headers) and headers[col_idx]:
                            row_data[headers[col_idx]] = cell.text.strip()
                        else:
                            row_data[f"Column_{col_idx+1}"] = cell.text.strip()
                    
                    if any(row_data.values()):  # Only add non-empty rows
                        table_content.append(row_data)
                
                if table_content:
                    tables_data.append(table_content)
            
            # Return appropriate format based on content
            if tables_data and not text_content:
                # Only tables
                if len(tables_data) == 1:
                    return self._clean_data_for_json(tables_data[0])
                else:
                    return self._clean_data_for_json(tables_data)
            elif text_content and not tables_data:
                # Only text
                return {"content": " ".join(text_content)}
            elif text_content and tables_data:
                # Both text and tables
                return {
                    "text": " ".join(text_content),
                    "tables": self._clean_data_for_json(tables_data if len(tables_data) > 1 else tables_data[0])
                }
            else:
                return {"content": "No content found in document"}
                
        except Exception as e:
            print(f"Error processing Word file: {e}")
            return None
    
    def process_file_with_supabase(self, file_path, file_type):
        """Process file and automatically insert into appropriate Supabase table"""
        # First extract the data
        data = self.process_file(file_path, file_type)
        
        if not data:
            print("❌ Failed to extract data from file")
            return None
        
        # Skip Supabase processing for Word documents (usually text, not tabular data)
        if file_type == 'word' and isinstance(data, dict) and 'content' in data:
            print("📄 Word document contains text content - not suitable for database insertion")
            return data
        
        # Handle multi-sheet Excel files
        if isinstance(data, dict) and file_type == 'excel':
            print("📊 Multi-sheet Excel detected - processing each sheet separately")
            results = {}
            
            for sheet_name, sheet_data in data.items():
                print(f"\n--- Processing Sheet: {sheet_name} ---")
                result = self._process_sheet_data(sheet_data, f"{file_path}_{sheet_name}")
                results[sheet_name] = result
            
            return results
        
        # Handle single data structure (CSV, single Excel sheet, Word tables)
        elif isinstance(data, list):
            return self._process_sheet_data(data, file_path)
        
        else:
            print("❌ Unsupported data structure for database insertion")
            return data
    
    def _process_sheet_data(self, data, source_name):
        """Process a single sheet/table of data for Supabase insertion"""
        if not data or not isinstance(data, list):
            if self.verbose:
                print(f"❌ No valid data to process from {source_name}")
            return data
        
        if self.verbose:
            print(f"📊 Analyzing data structure ({len(data)} records)...")
        
        # Analyze data structure
        data_structure = self.schema_mapper.analyze_data_structure(data)
        if self.verbose:
            print(f"   Columns: {', '.join(data_structure['columns'])}")
        
        # Find matching table
        table_name, confidence = self.schema_mapper.find_matching_table(data_structure)
        
        # Always display available tables and best match score
        available_tables = ", ".join(sorted(self.schema_mapper.table_schemas.keys()))
        print(f"📋 Available tables: {available_tables}")
        
        if not table_name:
            print(f"❌ No matching table schema found (best match confidence: {confidence:.2f})")
            
            # Show best match scores for all tables
            print("\n📊 Match scores for all tables:")
            all_scores = {}
            for t_name in self.schema_mapper.table_schemas.keys():
                score = self.schema_mapper._calculate_match_score(data_structure["columns"], data_structure, self.schema_mapper.table_schemas[t_name])
                all_scores[t_name] = score
                
            # Sort by score descending
            for t_name, score in sorted(all_scores.items(), key=lambda x: x[1], reverse=True):
                print(f"   • {t_name}: {score:.2f}" + (" ✓" if score >= 0.6 else " ✗"))
                
            return data
        
        if self.verbose:
            print(f"🎯 Matched to table: '{table_name}' (confidence: {confidence:.2f})")
        else:
            print(f"🎯 Matched to table: '{table_name}' (confidence: {confidence:.2f})")
        
        # DEBUG: Print sample data before validation
        if table_name in ['sms_header', 'tower_dumps']:
            print(f"\n🔍 DEBUG - Sample data before validation for {table_name}:")
            for i, record in enumerate(data[:3]):
                print(f"Record {i+1}:")
                for key, value in record.items():
                    print(f"  {key}: {type(value).__name__} = {value}")
        
        # Validate data for the matched table
        is_valid, validation_errors = self.schema_mapper.validate_data_for_table(data, table_name)
        
        if not is_valid:
            if self.verbose:
                print(f"❌ Invalid Schema - Data doesn't match table '{table_name}':")
                for error in validation_errors:
                    print(f"   • {error}")
            else:
                print(f"❌ Invalid Schema - Data doesn't match table '{table_name}'")
                for error in validation_errors:
                    print(f"   • {error}")
            
            # DEBUG: For sms_header or tower_dumps, try to fix the data directly here
            if table_name in ['sms_header', 'tower_dumps']:
                print(f"\n🔧 DEBUG - Attempting to fix data types for {table_name}...")
                fixed_data = []
                for record in data:
                    fixed_record = {}
                    for key, value in record.items():
                        # Handle specific data type conversions based on table
                        if table_name == 'sms_header':
                            # Convert all values to strings for sms_header
                            fixed_record[key] = str(value) if value is not None else None
                        elif table_name == 'tower_dumps':
                            # Get expected type for this column if it exists in schema
                            schema = self.schema_mapper.table_schemas[table_name]
                            column_mapping = self.schema_mapper._get_column_mapping([key], schema)
                            mapped_key = column_mapping.get(key, key)
                            expected_type = schema["column_types"].get(mapped_key)
                            
                            if value is None:
                                fixed_record[key] = None
                            elif expected_type == 'bigint' and (isinstance(value, (int, float)) or (isinstance(value, str) and value.isdigit())):
                                # Convert to int for bigint fields
                                try:
                                    fixed_record[key] = int(float(value))
                                except:
                                    fixed_record[key] = str(value)
                            elif expected_type == 'integer' and (isinstance(value, (int, float)) or (isinstance(value, str) and value.isdigit())):
                                # Convert to int for integer fields
                                try:
                                    fixed_record[key] = int(float(value))
                                except:
                                    fixed_record[key] = str(value)
                            elif expected_type == 'numeric' and isinstance(value, (int, float, str)):
                                # Convert to float for numeric fields
                                try:
                                    fixed_record[key] = float(value)
                                except:
                                    fixed_record[key] = str(value)
                            elif expected_type == 'text' or expected_type is None:
                                # Convert to string for text fields or unknown types
                                fixed_record[key] = str(value)
                            else:
                                # Keep original value for other types
                                fixed_record[key] = value
                        else:
                            # Default behavior for other tables
                            fixed_record[key] = value
                    fixed_data.append(fixed_record)
                
                # Print sample of fixed data
                print(f"\n🔍 DEBUG - Sample of fixed data for {table_name}:")
                for i, record in enumerate(fixed_data[:3]):
                    print(f"Record {i+1}:")
                    for key, value in record.items():
                        print(f"  {key}: {type(value).__name__} = {value}")
                
                # Try validation again
                is_valid_fixed, validation_errors_fixed = self.schema_mapper.validate_data_for_table(fixed_data, table_name)
                print(f"Fixed data validation: {'✅ Success' if is_valid_fixed else '❌ Failed'}")
                
                if is_valid_fixed:
                    print("Proceeding with fixed data...")
                    data = fixed_data
                    is_valid = True
                    validation_errors = []
                else:
                    print("Fixing failed. Errors:")
                    for error in validation_errors_fixed:
                        print(f"   • {error}")
                    return data
            else:
                return data
        
        # Normalize column names to match database schema
        normalized_data = self.schema_mapper.normalize_column_names(data, table_name)
        
        # DEBUG: Print sample data after normalization for monitored tables
        if table_name in ['sms_header', 'tower_dumps']:
            print(f"\n🔍 DEBUG - Sample data after normalization for {table_name}:")
            for i, record in enumerate(normalized_data[:3]):
                print(f"Record {i+1}:")
                for key, value in record.items():
                    print(f"  {key}: {type(value).__name__} = {value}")
        
        # Apply minimal data cleaning for JSON-to-database compatibility
        cleaned_data = self._clean_data_for_database(normalized_data, table_name)
        
        # DEBUG: Print sample data after cleaning for monitored tables
        if table_name in ['sms_header', 'tower_dumps']:
            print(f"\n🔍 DEBUG - Sample data after cleaning for {table_name}:")
            for i, record in enumerate(cleaned_data[:3]):
                print(f"Record {i+1}:")
                for key, value in record.items():
                    print(f"  {key}: {type(value).__name__} = {value}")
        
        # Insert data directly
        success, message, count = self.supabase_handler.insert_data(table_name, cleaned_data)
        
        if success:
            print(f"✅ SUCCESS: {count} records inserted into '{table_name}' table")
            return {
                "status": "success",
                "table": table_name,
                "records_inserted": count,
                "data": cleaned_data
            }
        else:
            print(f"❌ FAILED: Unable to insert data into '{table_name}' - {message}")
            return {
                "status": "failed",
                "table": table_name,
                "error": message,
                "data": cleaned_data
            }
    
    def _clean_data_for_database(self, data: List[Dict], table_name: str) -> List[Dict]:
        """Minimal data cleaning for database compatibility"""
        if table_name not in self.schema_mapper.table_schemas:
            return data
        
        schema = self.schema_mapper.table_schemas[table_name]
        column_types = schema.get("column_types", {})
        
        cleaned_data = []
        for record in data:
            cleaned_record = {}
            for key, value in record.items():
                if value is None:
                    cleaned_record[key] = value
                    continue
                
                expected_type = column_types.get(key)
                
                # Convert all values to text for sms_header table
                if table_name == 'sms_header':
                    cleaned_record[key] = str(value)
                # Special handling for tower_dumps table
                elif table_name == 'tower_dumps':
                    if expected_type == 'bigint' and (isinstance(value, (int, float)) or (isinstance(value, str) and value.replace('.', '', 1).isdigit())):
                        # Convert to int for bigint fields, handling possible float strings
                        try:
                            cleaned_record[key] = int(float(value))
                        except:
                            cleaned_record[key] = str(value)
                    elif expected_type == 'integer' and (isinstance(value, (int, float)) or (isinstance(value, str) and value.replace('.', '', 1).isdigit())):
                        # Convert to int for integer fields, handling possible float strings
                        try:
                            cleaned_record[key] = int(float(value))
                        except:
                            cleaned_record[key] = str(value)
                    elif expected_type == 'text' or key in ['b_party', 'call_type', 'first_cell_id_a', 'last_cell_id_a', 'first_cell_id_a_address', 'roaming_a']:
                        # Always convert these fields to string
                        cleaned_record[key] = str(value)
                    else:
                        # Keep original value for other types
                        cleaned_record[key] = value
                # Special handling for IPDR table
                elif table_name == 'ipdr':
                    if key in ['landline_msidn_mdn_leased_circuit_id', 'user_id', 'source_ip_address', 'translated_ip_address', 'destination_ip_address'] and expected_type == 'text':
                        # Convert these fields to string
                        cleaned_record[key] = str(value)
                    elif expected_type == 'bigint' and (isinstance(value, (int, float)) or (isinstance(value, str) and value.replace('.', '', 1).isdigit())):
                        # Convert to int for bigint fields, handling possible float strings
                        try:
                            cleaned_record[key] = int(float(value))
                        except:
                            cleaned_record[key] = str(value)
                    elif expected_type == 'integer' and (isinstance(value, (int, float)) or (isinstance(value, str) and value.replace('.', '', 1).isdigit())):
                        # Convert to int for integer fields, handling possible float strings
                        try:
                            cleaned_record[key] = int(float(value))
                        except:
                            cleaned_record[key] = str(value)
                    elif expected_type == 'text':
                        # Always convert text fields to string
                        cleaned_record[key] = str(value)
                    else:
                        # Keep original value for other types
                        cleaned_record[key] = value
                # For other tables, handle type conversion as before
                elif expected_type in ["integer", "bigint"] and isinstance(value, float) and value.is_integer():
                    # Convert float to int only if it's a whole number (e.g., 3478.0 → 3478)
                    cleaned_record[key] = int(value)
                else:
                    cleaned_record[key] = value
            
            cleaned_data.append(cleaned_record)
        
        return cleaned_data
    
    def process_file(self, file_path, file_type):
        """Main method to process any supported file type (original method)"""
        if file_type not in self.supported_types:
            print(f"Unsupported file type: {file_type}")
            return None
        
        # Call appropriate processor
        if file_type == 'csv':
            return self.process_csv(file_path)
        elif file_type == 'excel':
            return self.process_excel(file_path)
        elif file_type == 'word':
            return self.process_word(file_path)
        
        return None

# Test function for standalone use
def main():
    """Test the file processor directly"""
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        processor = FileProcessor()
        
        # Detect file type from extension
        ext = Path(file_path).suffix.lower()
        type_map = {
            '.csv': 'csv',
            '.xlsx': 'excel',
            '.xls': 'excel',
            '.docx': 'word',
            '.doc': 'word'
        }
        
        file_type = type_map.get(ext)
        if file_type:
            result = processor.process_file_with_supabase(file_path, file_type)
            if result:
                # Only print JSON if it's not a status result
                if isinstance(result, dict) and 'status' in result:
                    pass  # Status already printed by the processor
                else:
                    print(json.dumps(result, indent=2, ensure_ascii=False))
        else:
            print(f"Unsupported file type: {ext}")
    else:
        print("Usage: python files.py <file_path>")

if __name__ == "__main__":
    main() 